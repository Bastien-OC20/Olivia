"""
Production de documents Word à l'identité de l'établissement.

Règle d'architecture, non négociable : **le modèle de langage produit le CONTENU,
ce module produit la MISE EN FORME.** On ne demande jamais au modèle d'émettre du
XML ni du Markdown de mise en page — un modèle 7B n'en est pas capable de façon
fiable. On reçoit ici une structure simple (titre, objet, paragraphes, sous-titres,
listes, tableaux) et ce sont les styles Word du modèle de l'établissement qui sont
appliqués par le code.

Deux entrées possibles :
  - `contenu` : une liste de blocs déjà structurés (`paragraphe`, `sous_titre`,
    `liste`, `tableau`) — c'est le format cible ;
  - `texte` : la réponse d'Olivia telle qu'elle s'affiche dans la conversation.
    `structurer_texte()` la convertit en blocs. Ce chemin ne coûte AUCUN appel
    supplémentaire au modèle : le contenu a déjà été produit, on ne fait que le
    mettre en forme. C'est ce qui permet le « en un geste » demandé par l'interface.

Quatre types de documents partagent la même identité : circulaire, courrier aux
familles, convocation, compte rendu. Ce qui les distingue (bloc destinataire, lieu
et date, objet, formule de politesse, participants) est décrit dans `PROFILS` en
tête de fichier, pas dispersé dans la logique : ces choix sont des usages
d'établissement, l'utilisatrice doit pouvoir les corriger sans lire le code.

Dégradation propre : si le modèle de l'établissement est absent ou illisible, le
document est tout de même produit — sans le logo ni l'en-tête — et un
avertissement explicite est renvoyé à l'interface.
"""
import re
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from . import docmodele

# ---------- Textes d'usage, modifiables ----------
# Formules retenues d'après les usages d'un établissement scolaire français.
# Elles sont ici, en clair et en un seul endroit, PARCE QUE l'utilisatrice les
# corrigera : chacune est aussi surchargeable requête par requête (champs
# `appel`, `formule_politesse`, `lieu`, `signature` de la demande).
TEXTES = {
    "appel": "Madame, Monsieur,",
    "formule_politesse": "Je vous prie d'agréer, Madame, Monsieur, l'expression de mes "
                         "salutations distinguées.",
    "lieu": "Marseille",
    "prefixe_objet": "Objet : ",
    "prefixe_participants": "Participants : ",
    "prefixe_date_reunion": "Date : ",
    "signature": "",
}

# ---------- Profils de mise en page ----------
# Un profil décrit CE QUI COMPOSE le document, pas comment on l'écrit.
PROFILS = {
    "circulaire": {
        "libelle": "Circulaire",
        "titre_defaut": "Circulaire",
        "destinataire": False,
        "lieu_date": False,
        "objet": False,
        "appel": False,
        "politesse": False,
        "participants": False,
        "bandeau_sous_titres": True,
    },
    "courrier": {
        "libelle": "Courrier aux familles",
        "titre_defaut": "Courrier aux familles",
        "destinataire": True,
        "lieu_date": True,
        "objet": True,
        "appel": True,
        "politesse": True,
        "participants": False,
        "bandeau_sous_titres": False,
    },
    "convocation": {
        "libelle": "Convocation",
        "titre_defaut": "Convocation",
        "destinataire": True,
        "lieu_date": True,
        "objet": True,
        "appel": True,
        "politesse": True,
        "participants": False,
        "bandeau_sous_titres": False,
    },
    "compte_rendu": {
        "libelle": "Compte rendu",
        "titre_defaut": "Compte rendu",
        "destinataire": False,
        "lieu_date": True,
        "objet": False,
        "appel": False,
        "politesse": False,
        "participants": True,
        "bandeau_sous_titres": True,
    },
}

TYPES = tuple(PROFILS)

# ---------- Styles du modèle ----------
# Noms tels qu'ils existent dans la circulaire de l'établissement. `_style()`
# retombe sur "Normal" si l'un d'eux manque (document sans modèle).
STYLE_TITRE = "Heading 1"
STYLE_SOUS_TITRE = "Heading 5"
STYLE_CORPS = "Body Text"
STYLE_LISTE = "List Paragraph"

# Gris du bandeau des intertitres, relevé sur la circulaire de rentrée.
GRIS_BANDEAU = "BEBEBE"

# Ordre imposé par le schéma OOXML pour les enfants de w:pPr : `w:shd` doit
# précéder tous ces éléments. Sans cela Word refuse d'ouvrir le document.
_APRES_SHD = (
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap", "w:overflowPunct",
    "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
    "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
    "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr",
    "w:sectPr", "w:pPrChange",
)

MOIS = ("janvier", "février", "mars", "avril", "mai", "juin", "juillet",
        "août", "septembre", "octobre", "novembre", "décembre")

# Bornes : un document reste un document. Elles évitent qu'une réponse
# aberrante du modèle produise un fichier de plusieurs milliers de pages.
MAX_BLOCS = 400
MAX_ITEMS_LISTE = 100
MAX_LIGNES_TABLEAU = 60
MAX_COLONNES_TABLEAU = 12
MAX_CARACTERES_BLOC = 5000


# ---------- Structuration du texte ----------
_RE_THINK = re.compile(r"<think>.*?</think>", re.S | re.I)
_RE_TITRE_MD = re.compile(r"^(#{1,6})\s+(.*)$")
_RE_PUCE = re.compile(r"^\s*[-*•·]\s+(.*)$")
_RE_NUMERO = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_RE_SEPARATEUR = re.compile(r"^\s*([-*_])\1{2,}\s*$")
_RE_LIGNE_TABLEAU = re.compile(r"^\s*\|.*\|\s*$")
_RE_SEPARATEUR_TABLEAU = re.compile(r"^\s*\|[\s:|-]+\|\s*$")
_RE_GRAS = re.compile(r"\*\*(.+?)\*\*")


def _cellules(ligne: str) -> list:
    return [c.strip() for c in ligne.strip().strip("|").split("|")]


def structurer_texte(texte: str) -> list:
    """Convertit la réponse d'Olivia (texte courant, Markdown léger) en blocs.

    Aucun appel au modèle : le contenu existe déjà, on ne fait que le lire. Le
    bloc de raisonnement `<think>…</think>` des modèles qui en émettent est
    retiré — il n'a rien à faire dans un courrier officiel.
    """
    brut = _RE_THINK.sub("", texte or "")
    brut = re.sub(r"<think>.*", "", brut, flags=re.S | re.I)
    lignes = brut.replace("\r\n", "\n").split("\n")
    blocs: list = []
    i = 0
    while i < len(lignes) and len(blocs) < MAX_BLOCS:
        ligne = lignes[i]
        nue = ligne.strip()
        if not nue or _RE_SEPARATEUR.match(nue):
            i += 1
            continue

        # Tableau Markdown : au moins un en-tête et une ligne de séparation.
        if (_RE_LIGNE_TABLEAU.match(nue) and i + 1 < len(lignes)
                and _RE_SEPARATEUR_TABLEAU.match(lignes[i + 1])):
            lignes_tab = [_cellules(nue)]
            i += 2
            while i < len(lignes) and _RE_LIGNE_TABLEAU.match(lignes[i].strip()):
                lignes_tab.append(_cellules(lignes[i].strip()))
                i += 1
                if len(lignes_tab) >= MAX_LIGNES_TABLEAU:
                    break
            blocs.append({"type": "tableau", "lignes": lignes_tab})
            continue

        m = _RE_TITRE_MD.match(nue)
        if m:
            blocs.append({"type": "sous_titre", "texte": m.group(2).strip(),
                          "niveau": len(m.group(1))})
            i += 1
            continue

        m = _RE_PUCE.match(ligne) or _RE_NUMERO.match(ligne)
        if m:
            items = []
            while i < len(lignes) and len(items) < MAX_ITEMS_LISTE:
                m2 = _RE_PUCE.match(lignes[i]) or _RE_NUMERO.match(lignes[i])
                if not m2:
                    break
                items.append(m2.group(1).strip())
                i += 1
            blocs.append({"type": "liste", "items": items})
            continue

        # Paragraphe : on recolle les lignes coupées à la main. Un texte collé
        # depuis un courriel arrive souvent replié à 80 colonnes ; une nouvelle
        # phrase, elle, commence par une majuscule.
        morceaux = [nue]
        i += 1
        while i < len(lignes):
            suite = lignes[i].strip()
            if not suite or _RE_SEPARATEUR.match(suite) or _RE_TITRE_MD.match(suite):
                break
            if _RE_PUCE.match(lignes[i]) or _RE_NUMERO.match(lignes[i]):
                break
            if _RE_LIGNE_TABLEAU.match(suite) or not suite[:1].islower():
                break
            morceaux.append(suite)
            i += 1
        blocs.append({"type": "paragraphe", "texte": " ".join(morceaux)})
    return blocs


def titre_probable(texte: str) -> str:
    """Titre déduit d'un texte : premier titre Markdown, sinon première ligne."""
    for bloc in structurer_texte(texte)[:3]:
        if bloc["type"] == "sous_titre":
            return bloc["texte"]
    for bloc in structurer_texte(texte)[:1]:
        if bloc["type"] == "paragraphe":
            mots = bloc["texte"].split()
            return " ".join(mots[:12])
    return ""


# ---------- Outils de mise en forme ----------
def _style(doc, nom: str):
    """Style Word portant ce nom dans ce document, sinon « Normal ».

    On renvoie l'OBJET style et non son nom : la recherche par nom de
    python-docx traduit les noms intégrés (« Heading 1 » → « heading 1 ») et
    échoue sur les documents produits par Word/LibreOffice, qui écrivent
    « Heading 1 » tel quel. On compare donc aussi l'identifiant de style.
    """
    sans_espace = nom.replace(" ", "")
    for style in doc.styles:
        if style.name == nom or style.style_id == sans_espace:
            return style
    try:
        return doc.styles[nom]
    except KeyError:
        return doc.styles["Normal"]


def _bandeau_gris(paragraphe, couleur: str = GRIS_BANDEAU) -> None:
    """Fond gris pleine largeur — signature visuelle des intertitres du lycée."""
    pPr = paragraphe._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), couleur)
    pPr.insert_element_before(shd, *_APRES_SHD)


def _sans_indentation(paragraphe) -> None:
    """Annule les retraits hérités d'un style (le titre du modèle en porte)."""
    paragraphe.paragraph_format.left_indent = 0
    paragraphe.paragraph_format.right_indent = 0
    paragraphe.paragraph_format.first_line_indent = 0


def _ecrire_riche(paragraphe, texte: str) -> None:
    """Écrit le texte en gérant `**gras**`. Le reste est écrit tel quel.

    C'est bien du CONTENU balisé légèrement par le modèle, pas de la mise en
    page : le gras d'insistance est converti en gras Word, rien d'autre.
    """
    texte = (texte or "")[:MAX_CARACTERES_BLOC]
    position = 0
    for m in _RE_GRAS.finditer(texte):
        if m.start() > position:
            paragraphe.add_run(texte[position:m.start()])
        paragraphe.add_run(m.group(1)).bold = True
        position = m.end()
    reste = texte[position:]
    if reste or not paragraphe.runs:
        paragraphe.add_run(reste)


def _numero_puce(doc) -> Optional[int]:
    """numId d'une liste à puces définie dans le document, sinon None.

    Le modèle de l'établissement en définit une ; le document de repli (sans
    modèle) peut ne pas en avoir — on retombe alors sur une puce écrite.
    """
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return None
    abstraits = set()
    for abstrait in numbering.findall(qn("w:abstractNum")):
        niveau = abstrait.find(qn("w:lvl"))
        if niveau is None:
            continue
        fmt = niveau.find(qn("w:numFmt"))
        if fmt is not None and fmt.get(qn("w:val")) == "bullet":
            abstraits.add(abstrait.get(qn("w:abstractNumId")))
    for num in numbering.findall(qn("w:num")):
        ref = num.find(qn("w:abstractNumId"))
        if ref is not None and ref.get(qn("w:val")) in abstraits:
            try:
                return int(num.get(qn("w:numId")))
            except (TypeError, ValueError):
                continue
    return None


def _appliquer_puce(paragraphe, num_id: int) -> None:
    """Rattache le paragraphe à la liste à puces définie dans le document."""
    numPr = paragraphe._p.get_or_add_pPr().get_or_add_numPr()
    numPr.get_or_add_ilvl().val = 0
    numPr.get_or_add_numId().val = num_id


def _habiller_tableau(table) -> None:
    """Filets fins et marges de cellule : le modèle ne définit aucun style bordé,
    et son style « Table Normal » colle le texte contre le trait."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for cote in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{cote}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)
    marges = OxmlElement("w:tblCellMar")
    for cote, largeur in (("top", 40), ("left", 80), ("bottom", 40), ("right", 80)):
        el = OxmlElement(f"w:{cote}")
        el.set(qn("w:w"), str(largeur))
        el.set(qn("w:type"), "dxa")
        marges.append(el)
    tblPr.append(marges)


def date_en_francais(valeur: str = "") -> str:
    """« 28 juillet 2026 ». Accepte une date ISO, sinon renvoie la valeur telle quelle."""
    brut = (valeur or "").strip()
    if not brut:
        jour = date.today()
    else:
        m = re.match(r"^(\d{4})-(\d{2})-(\d{2})", brut)
        if not m:
            return brut
        try:
            jour = date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except ValueError:
            return brut
    premier = "1er" if jour.day == 1 else str(jour.day)
    return f"{premier} {MOIS[jour.month - 1]} {jour.year}"


# ---------- Assemblage du document ----------
def _nouveau_document(chemin_modele: Optional[Path]):
    """Document fondé sur le modèle de l'établissement, ou vierge à défaut."""
    if chemin_modele and Path(chemin_modele).is_file():
        try:
            return Document(str(chemin_modele)), ""
        except Exception as e:
            return Document(), (f"Le modèle de l'établissement est illisible ({e}) : "
                                "le document a été créé sans le logo ni l'en-tête.")
    return Document(), ("Le modèle de l'établissement est introuvable : le document a "
                        "été créé sans le logo ni l'en-tête du lycée.")


def _ajouter_titre(doc, texte: str) -> None:
    p = doc.add_paragraph(style=_style(doc, STYLE_TITRE))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _ecrire_riche(p, texte)


def _ajouter_sous_titre(doc, texte: str, bandeau: bool) -> None:
    p = doc.add_paragraph(style=_style(doc, STYLE_SOUS_TITRE))
    _sans_indentation(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bandeau else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    _ecrire_riche(p, texte)
    if bandeau:
        _bandeau_gris(p)


def _ajouter_paragraphe(doc, texte: str, alignement=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        gras: bool = False):
    p = doc.add_paragraph(style=_style(doc, STYLE_CORPS))
    _sans_indentation(p)
    p.alignment = alignement
    p.paragraph_format.space_after = Pt(6)
    _ecrire_riche(p, texte)
    if gras:
        for run in p.runs:
            run.bold = True
    return p


def _ajouter_liste(doc, items: list, num_id: Optional[int]) -> None:
    for item in items[:MAX_ITEMS_LISTE]:
        p = doc.add_paragraph(style=_style(doc, STYLE_LISTE))
        p.paragraph_format.space_after = Pt(2)
        if num_id is None:
            _sans_indentation(p)
            _ecrire_riche(p, f"• {item}")
        else:
            _ecrire_riche(p, str(item))
            _appliquer_puce(p, num_id)


def _ajouter_tableau(doc, lignes: list) -> None:
    lignes = [ligne[:MAX_COLONNES_TABLEAU] for ligne in lignes[:MAX_LIGNES_TABLEAU]]
    lignes = [ligne for ligne in lignes if ligne]
    if not lignes:
        return
    colonnes = max(len(ligne) for ligne in lignes)
    table = doc.add_table(rows=len(lignes), cols=colonnes)
    _habiller_tableau(table)
    for i, ligne in enumerate(lignes):
        for j in range(colonnes):
            cellule = table.cell(i, j)
            para = cellule.paragraphs[0]
            para.style = _style(doc, STYLE_CORPS)
            _ecrire_riche(para, ligne[j] if j < len(ligne) else "")
            if i == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph(style=_style(doc, STYLE_CORPS))


def _ajouter_contenu(doc, blocs: list, bandeau: bool) -> None:
    num_id = _numero_puce(doc)
    for bloc in blocs[:MAX_BLOCS]:
        genre = (bloc.get("type") or "paragraphe").strip()
        if genre == "sous_titre":
            _ajouter_sous_titre(doc, bloc.get("texte", ""), bandeau)
        elif genre == "liste":
            _ajouter_liste(doc, list(bloc.get("items") or []), num_id)
        elif genre == "tableau":
            _ajouter_tableau(doc, list(bloc.get("lignes") or []))
        else:
            _ajouter_paragraphe(doc, bloc.get("texte", ""))


def _valeur(demande: dict, cle: str, defaut: str = "") -> str:
    return str(demande.get(cle) or defaut).strip()


def _bloc_utile(bloc) -> bool:
    """Écarte les blocs vides : un modèle 7B en produit régulièrement."""
    if not isinstance(bloc, dict):
        return False
    genre = (bloc.get("type") or "paragraphe").strip()
    if genre == "liste":
        return any(str(x).strip() for x in (bloc.get("items") or []))
    if genre == "tableau":
        return any(any(str(c).strip() for c in ligne) for ligne in (bloc.get("lignes") or []))
    return bool(str(bloc.get("texte") or "").strip())


# Ouvertures et clôtures qu'un modèle ajoute spontanément quand on lui demande
# « rédige une convocation » : il produit une lettre ENTIÈRE, alors que le
# profil du document fournit déjà l'objet, l'appel, la formule de politesse et
# la signature. Sans ce nettoyage, le document final les porte EN DOUBLE —
# constaté en conditions réelles avec Mistral Nemo.
_RE_OBJET = re.compile(r"^objet\s*:", re.I)
_RE_APPEL = re.compile(r"^(madame|monsieur|chers?|cher)\b[^.!?]{0,60},\s*$", re.I)
# Motif NON ancré : la formule arrive souvent en fin de phrase (« En espérant
# vous compter parmi nous, veuillez agréer… »), pas en début de paragraphe.
_RE_CLOTURE = re.compile(
    r"(cordialement|sinc[èe]rement|salutations distingu|salutations respectueuses"
    r"|je vous prie d.agr[ée]er|veuillez agr[ée]er|dans l.attente de vous"
    r"|esp[ée]r\w* vous compter|en esp[ée]rant vous|vous remerciant"
    r"|restons à votre disposition|nous vous remercions de votre)", re.I)
# Lignes de signature : « Le Chef d'établissement », « Professeur principal de
# la 2nde B », ou une mention laissée en attente « [Nom du professeur] ».
# L'article est facultatif : le modèle l'omet une fois sur deux.
_RE_SIGNATURE = re.compile(
    r"^([\[\(<].{0,60}[\]\)>]"
    r"|((le|la|les)\s+)?(chef|proviseur|proviseure|directeur|directrice"
    r"|principal|principale|professeur|professeure|cpe|secr[ée]taire"
    r"|direction|[ée]quipe p[ée]dagogique)[^.!?]{0,60})\s*$", re.I)

# Au-delà, on mordrait sur le propos de l'utilisatrice.
_MAX_RETRAITS_QUEUE = 5


def _texte_bloc(bloc) -> str:
    return str(bloc.get("texte") or "").strip() if isinstance(bloc, dict) else ""


def _degarnir(blocs: list, profil: dict) -> list:
    """Retire du corps ce que le document fournit déjà par ailleurs.

    N'intervient qu'en TÊTE et en QUEUE : un « Cordialement » au milieu d'un
    texte appartient au propos de l'utilisatrice, on n'y touche pas.
    """
    blocs = list(blocs)
    # Tête : objet et appel, seulement si le profil les pose lui-même.
    while blocs:
        t = _texte_bloc(blocs[0])
        if not t:
            break
        if (profil.get("objet") and _RE_OBJET.match(t)) or \
                (profil.get("appel") and _RE_APPEL.match(t)):
            blocs.pop(0)
            continue
        break
    # Queue : formule de politesse puis lignes de signature. On remonte depuis
    # la fin tant que ça ressemble à une clôture, sans jamais vider le document
    # ni dépasser _MAX_RETRAITS_QUEUE.
    if profil.get("politesse"):
        retires = 0
        while len(blocs) > 1 and retires < _MAX_RETRAITS_QUEUE:
            t = _texte_bloc(blocs[-1])
            if t and (_RE_CLOTURE.search(t) or _RE_SIGNATURE.match(t)):
                blocs.pop()
                retires += 1
                continue
            break
    return blocs


def generer(demande: dict, destination: Path,
            chemin_modele: Optional[Path] = None) -> dict:
    """Écrit le document Word demandé et renvoie un compte rendu.

    `demande` : type, titre, objet, texte OU contenu, destinataire, lieu, date,
    participants, appel, formule_politesse, signature.
    Lève `ValueError` si le type est inconnu ou si le contenu est vide.
    """
    genre = _valeur(demande, "type", "circulaire")
    profil = PROFILS.get(genre)
    if profil is None:
        raise ValueError(f"Type de document inconnu : {genre}")

    blocs = [b for b in (demande.get("contenu") or []) if _bloc_utile(b)]
    if not blocs:
        blocs = structurer_texte(_valeur(demande, "texte"))
    if not blocs:
        raise ValueError("Le document serait vide : aucun contenu fourni.")

    # Un modèle sollicité pour « une convocation » rédige la lettre entière :
    # on retire ce que le profil pose déjà, sinon tout figure en double.
    blocs = _degarnir(blocs, profil)
    if not blocs:
        raise ValueError("Le document serait vide : aucun contenu fourni.")

    if chemin_modele is None:
        chemin_modele = docmodele.chemin_modele()
    doc, avertissement = _nouveau_document(chemin_modele)

    titre = _valeur(demande, "titre") or profil["titre_defaut"]
    _ajouter_titre(doc, titre)
    doc.add_paragraph(style=_style(doc, STYLE_CORPS))

    if profil["destinataire"]:
        destinataire = _valeur(demande, "destinataire")
        for ligne in [x for x in destinataire.split("\n") if x.strip()][:8]:
            _ajouter_paragraphe(doc, ligne.strip(), WD_ALIGN_PARAGRAPH.RIGHT)

    if profil["lieu_date"]:
        lieu = _valeur(demande, "lieu", TEXTES["lieu"])
        jour = date_en_francais(_valeur(demande, "date"))
        if genre == "compte_rendu":
            _ajouter_paragraphe(doc, f"{TEXTES['prefixe_date_reunion']}{jour} — {lieu}",
                                WD_ALIGN_PARAGRAPH.LEFT)
        else:
            _ajouter_paragraphe(doc, f"{lieu}, le {jour}", WD_ALIGN_PARAGRAPH.RIGHT)

    if profil["participants"]:
        participants = [str(x).strip() for x in (demande.get("participants") or [])
                        if str(x).strip()]
        if participants:
            _ajouter_paragraphe(doc, TEXTES["prefixe_participants"].rstrip(),
                                WD_ALIGN_PARAGRAPH.LEFT, gras=True)
            _ajouter_liste(doc, participants, _numero_puce(doc))

    if profil["objet"]:
        objet = _valeur(demande, "objet")
        if objet:
            _ajouter_paragraphe(doc, f"{TEXTES['prefixe_objet']}{objet}",
                                WD_ALIGN_PARAGRAPH.LEFT, gras=True)

    if profil["appel"]:
        _ajouter_paragraphe(doc, _valeur(demande, "appel", TEXTES["appel"]),
                            WD_ALIGN_PARAGRAPH.LEFT)

    _ajouter_contenu(doc, blocs, profil["bandeau_sous_titres"])

    if profil["politesse"]:
        _ajouter_paragraphe(
            doc,
            _valeur(demande, "formule_politesse", TEXTES["formule_politesse"]),
        )

    signature = _valeur(demande, "signature", TEXTES["signature"])
    if signature:
        doc.add_paragraph(style=_style(doc, STYLE_CORPS))
        for ligne in [x for x in signature.split("\n") if x.strip()][:4]:
            _ajouter_paragraphe(doc, ligne.strip(), WD_ALIGN_PARAGRAPH.RIGHT)

    # RGPD : le fichier produit ne porte aucun nom d'utilisateur système.
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = titre

    destination = Path(destination)
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))
    return {
        "chemin": str(destination),
        "type": genre,
        "libelle": profil["libelle"],
        "titre": titre,
        "blocs": len(blocs),
        "avertissement": avertissement,
    }
