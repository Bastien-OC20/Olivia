"""
Prévisualisation simplifiée de fichiers pour l'UI, et extraction de texte brut
pour la recherche documentaire.

`preview()` renvoie une structure uniforme selon le type :
  {kind: "text",  content}                 → texte / code
  {kind: "table", columns, rows, truncated}→ CSV / XLSX
  {kind: "doc",   paragraphs, truncated}   → DOCX
  {kind: "image", url, name}               → images (rendu via /api/fs/download)
  {kind: "pdf",   url, name}               → PDF (rendu natif navigateur)
  {kind: "unsupported", name, url}         → binaire non prévisualisable

`extract_text()` est distincte : elle produit le texte brut INTÉGRAL d'un
document (Word, Excel, PDF, texte, CSV) pour permettre de le fouiller. Elle ne
lève jamais : un fichier illisible, protégé ou corrompu renvoie "" plutôt que
de faire échouer la recherche entière. Quand un PDF n'a pas de couche texte
(document passé au scanner) ou qu'il s'agit d'une image, elle passe le relais à
`ocr.py` — c'est le SEUL point d'entrée de l'extraction, de sorte que la
recherche documentaire en bénéficie sans rien changer de son côté.
`extract_text_meta()` rend en plus la provenance du texte (reconnu ou natif).

Aucune dépendance lourde n'est importée au niveau module : docx/openpyxl/pypdf
sont chargés paresseusement pour que l'app démarre même si elles manquent.
"""
from pathlib import Path

TEXT_EXT = {".txt", ".md", ".py", ".js", ".ts", ".vue", ".json", ".yaml",
            ".yml", ".html", ".css", ".log", ".sh", ".ini", ".cfg", ".xml"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".bmp"}

# Images dont on sait tirer du texte par reconnaissance de caractères. Le GIF
# animé et le SVG en sont écartés : le premier n'est jamais un document scanné,
# le second contient déjà son texte sous forme de balises.
OCR_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}

MAX_ROWS = 200
MAX_CHARS = 30_000
MAX_PARAGRAPHS = 800

# Extensions dont on sait extraire du texte : c'est le périmètre de la recherche.
# Volontairement centré sur les documents de bureau (Word, Excel, PDF) et non sur
# les extensions de développeur — le fonds réel est fait de courriers, comptes
# rendus, convocations et exports de logiciels scolaires.
OFFICE_EXT = {".docx", ".xlsx", ".xlsm", ".pdf", ".csv"}
SEARCHABLE_EXT = TEXT_EXT | OFFICE_EXT

# Bornes d'extraction : un document de secrétariat dépasse rarement ces valeurs,
# et au-delà on préfère un extrait au blocage de la recherche.
EXTRACT_MAX_CHARS = 400_000
EXTRACT_MAX_XLSX_ROWS = 5_000
EXTRACT_MAX_PDF_PAGES = 60

# Seuil de décision « ce PDF est un scan ». Un document passé au scanner n'a en
# principe aucune couche texte, mais il en porte parfois quelques miettes : un
# tampon d'en-tête ajouté par le copieur, un numéro de page, un nom de fichier
# incrusté par le logiciel de numérisation. Une vraie page de courrier compte
# plusieurs centaines de caractères ; en dessous de 24 caractères alphanumériques
# par page en moyenne, il n'y a pas de texte exploitable et l'OCR est justifié.
SEUIL_ALNUM_PAR_PAGE = 24


def read_text_file(path: Path) -> str:
    """Lit un fichier texte en devinant son encodage.

    `Path.read_text()` utiliserait l'encodage de la console Windows (cp1252) et
    transformerait « élèves » en « Ã©lÃ¨ves » dans un fichier UTF-8 — ce qui
    rendrait le mot introuvable à la recherche. On tente donc l'UTF-8 d'abord,
    puis les encodages Windows historiques (cp1252, qui n'échoue jamais en
    dernier recours sous sa variante latin-1).
    """
    data = path.read_bytes()
    for enc in ("utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def preview(path: Path, download_url: str) -> dict:
    ext = path.suffix.lower()

    if ext in IMAGE_EXT:
        return {"kind": "image", "url": download_url, "name": path.name}
    if ext == ".pdf":
        return {"kind": "pdf", "url": download_url, "name": path.name}
    if ext == ".csv":
        return _preview_csv(path)
    if ext in {".xlsx", ".xlsm"}:
        return _preview_xlsx(path)
    if ext == ".docx":
        return _preview_docx(path)
    if ext in TEXT_EXT:
        try:
            content = read_text_file(path)
        except Exception as e:
            return {"kind": "error", "detail": f"Lecture impossible : {e}"}
        truncated = len(content) > MAX_CHARS
        return {"kind": "text", "content": content[:MAX_CHARS], "truncated": truncated}

    return {"kind": "unsupported", "name": path.name, "url": download_url}


def _preview_csv(path: Path) -> dict:
    import csv
    rows = []
    truncated = False
    try:
        with open(path, newline="", encoding="utf-8", errors="ignore") as f:
            for i, row in enumerate(csv.reader(f)):
                if i >= MAX_ROWS + 1:
                    truncated = True
                    break
                rows.append([("" if c is None else str(c)) for c in row])
    except Exception as e:
        return {"kind": "error", "detail": f"Erreur CSV : {e}"}
    columns = rows[0] if rows else []
    return {"kind": "table", "columns": columns, "rows": rows[1:], "truncated": truncated}


def _preview_xlsx(path: Path) -> dict:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"kind": "error", "detail": "Module 'openpyxl' non installé (pip install openpyxl)."}
    try:
        wb = load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        rows = []
        truncated = False
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= MAX_ROWS + 1:
                truncated = True
                break
            rows.append([("" if c is None else str(c)) for c in row])
        wb.close()
    except Exception as e:
        return {"kind": "error", "detail": f"Erreur XLSX : {e}"}
    columns = rows[0] if rows else []
    return {"kind": "table", "columns": columns, "rows": rows[1:],
            "sheet": getattr(ws, "title", ""), "truncated": truncated}


def _preview_docx(path: Path) -> dict:
    try:
        import docx  # python-docx
    except ImportError:
        return {"kind": "error",
                "detail": "Module 'python-docx' non installé (pip install python-docx)."}
    try:
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
    except Exception as e:
        return {"kind": "error", "detail": f"Erreur DOCX : {e}"}
    truncated = len(paras) > MAX_PARAGRAPHS
    return {"kind": "doc", "paragraphs": paras[:MAX_PARAGRAPHS], "truncated": truncated}


# ---------- Extraction de texte brut (recherche documentaire) ----------
# `profile_id` est le premier paramètre des fonctions d'extraction : la
# reconnaissance de caractères est un RÉGLAGE (ocr_enabled, ocr_tesseract_path),
# et les réglages appartiennent désormais à une organisation. Il n'est utilisé que
# pour cela — le contenu extrait, lui, ne dépend jamais du profil.
def ocr_actif(profile_id: str) -> bool:
    """Vrai si cette organisation a activé la reconnaissance de caractères."""
    try:
        from . import ocr
        return ocr.est_active(profile_id)
    except Exception:
        return False


def est_cherchable(profile_id: str, ext: str) -> bool:
    """Une extension entre-t-elle dans le périmètre de la recherche ?

    Les images n'y entrent que si la reconnaissance de caractères est active :
    sinon, les ouvrir ne produirait rien et gaspillerait le budget de fichiers
    examinés d'une recherche.
    """
    ext = (ext or "").lower()
    if ext in SEARCHABLE_EXT:
        return True
    return ext in OCR_IMAGE_EXT and ocr_actif(profile_id)


def extract_text(profile_id: str, path: Path) -> str:
    """Texte brut intégral d'un document, pour la recherche."""
    return extract_text_meta(profile_id, path)[0]


def extract_text_meta(profile_id: str, path: Path) -> tuple[str, dict]:
    """Texte brut intégral d'un document + provenance de ce texte.

    Couvre le texte/code, le CSV, le Word (.docx), l'Excel (.xlsx/.xlsm), le PDF
    et — par reconnaissance de caractères — les PDF scannés et les images.
    Renvoie "" si le type n'est pas couvert, si la dépendance manque, ou si la
    lecture échoue — jamais d'exception : un seul fichier abîmé ne doit pas
    interrompre une recherche portant sur des milliers de fichiers.

    Le dictionnaire renvoyé vaut {"ocr": bool, "notice": str} : `ocr` dit que le
    texte vient d'une transcription automatique (l'interface doit le signaler,
    car la reconnaissance n'est jamais parfaite), `notice` porte l'information
    quand une borne a coupé le traitement.
    """
    ext = path.suffix.lower()
    vide: dict = {"ocr": False, "notice": ""}
    try:
        if ext == ".docx":
            return _text_docx(path)[:EXTRACT_MAX_CHARS], vide
        if ext in {".xlsx", ".xlsm"}:
            return _text_xlsx(path)[:EXTRACT_MAX_CHARS], vide
        if ext == ".pdf":
            texte, meta = _text_pdf_ou_ocr(profile_id, path)
            return texte[:EXTRACT_MAX_CHARS], meta
        if ext in OCR_IMAGE_EXT:
            texte, meta = _text_image_ocr(profile_id, path)
            return texte[:EXTRACT_MAX_CHARS], meta
        if ext == ".csv" or ext in TEXT_EXT:
            return read_text_file(path)[:EXTRACT_MAX_CHARS], vide
    except Exception:
        return "", vide
    return "", vide


def _semble_scanne(texte: str, pages: int) -> bool:
    """Le PDF est-il dépourvu de couche texte exploitable ? (voir SEUIL_ALNUM_PAR_PAGE)"""
    alnum = sum(1 for c in texte if c.isalnum())
    return alnum < SEUIL_ALNUM_PAR_PAGE * max(1, pages)


def _text_pdf_ou_ocr(profile_id: str, path: Path) -> tuple[str, dict]:
    """Couche texte du PDF ; à défaut, reconnaissance de caractères.

    L'OCR n'est tenté QUE si la couche texte est vide ou quasi vide : un PDF
    ordinaire n'est donc jamais ralenti. Le peu de texte natif trouvé est
    conservé et placé devant le texte reconnu.
    """
    natif, pages = _text_pdf(path)
    if not _semble_scanne(natif, pages) or not ocr_actif(profile_id):
        return natif, {"ocr": False, "notice": ""}
    from . import ocr
    resultat = ocr.texte_pdf(profile_id, path)
    if resultat is None or not resultat.texte.strip():
        return natif, {"ocr": False, "notice": resultat.notice if resultat else ""}
    parts = [p for p in (natif.strip(), resultat.texte) if p]
    return "\n".join(parts), {"ocr": True, "notice": resultat.notice}


def _text_image_ocr(profile_id: str, path: Path) -> tuple[str, dict]:
    """Texte d'une image (circulaire photographiée, capture d'écran, fax)."""
    if not ocr_actif(profile_id):
        return "", {"ocr": False, "notice": ""}
    from . import ocr
    resultat = ocr.texte_image(profile_id, path)
    if resultat is None or not resultat.texte.strip():
        return "", {"ocr": False, "notice": resultat.notice if resultat else ""}
    return resultat.texte, {"ocr": True, "notice": resultat.notice}


def _text_docx(path: Path) -> str:
    """Paragraphes + cellules de tableaux (les convocations tiennent souvent
    dans un tableau Word, qui serait invisible en ne lisant que les paragraphes)."""
    try:
        import docx  # python-docx
    except ImportError:
        return ""
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _text_xlsx(path: Path) -> str:
    """Toutes les feuilles, une ligne de texte par ligne de tableur."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ""
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for ws in wb.worksheets:
            parts.append(str(ws.title))
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= EXTRACT_MAX_XLSX_ROWS:
                    break
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append("\t".join(cells))
    finally:
        wb.close()
    return "\n".join(parts)


def _text_pdf(path: Path) -> tuple[str, int]:
    """Couche texte du PDF (pypdf) et nombre de pages lues.

    Une page illisible est sautée. Le nombre de pages sert ensuite à décider si
    le document est un scan (voir `_semble_scanne`)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", 0
    reader = PdfReader(str(path))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")  # PDF « protégé » sans mot de passe réel
        except Exception:
            return "", 0
    parts: list[str] = []
    lues = 0
    for i, page in enumerate(reader.pages):
        if i >= EXTRACT_MAX_PDF_PAGES:
            break
        lues += 1
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(parts), lues
