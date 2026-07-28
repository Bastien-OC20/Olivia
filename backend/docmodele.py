"""
Fabrication du modèle Word de l'établissement à partir d'un document réel.

Pourquoi ce module existe : l'identité visuelle du lycée (l'ours — les deux
logos en tête de page — et le bandeau d'adresse en pied) n'est PAS reconstituable
en partant d'un document vierge. Elle tient à des images ancrées au millimètre,
à trois sections, à des styles Word (Calibri 11, Heading 1/2/4/5, Body Text) et à
des marges très particulières. La seule façon fiable de la conserver est donc de
REPARTIR du document réel et d'en retirer la rédaction.

Ce que produit `construire_modele()` :
  - un .docx d'une seule section, VIDE de tout texte ;
  - l'ours en en-tête de page, le bandeau d'adresse en pied, les deux ancrés
    à la page (donc rigoureusement au même endroit sur toutes les pages) ;
  - les styles, la numérotation, le thème et les polices du document source ;
  - aucune métadonnée d'auteur (RGPD).

Constat fait sur la circulaire de rentrée 2026-2027, qui corrige une idée reçue :
l'ours n'est effectivement PAS en en-tête sur la première page (section 1, marge
haute 0,64 cm, images ancrées dans le corps), mais il l'est sur les sections
suivantes (marge haute 2,61 cm, MÊMES images, ancrées à la page). Le modèle
retient cette seconde forme : elle donne exactement le même rendu, sur toutes
les pages, sans image flottante dans le corps — que la rédaction ferait dériver.
Si un jour la charte change et que le nouveau document ne porte l'ours que dans
le corps, le repli plus bas (`_ENTETE_DANS_LE_CORPS`) le conserve tel quel.

Fabrication rejouable : par la route `POST /api/documents/modele`, ou en ligne
de commande depuis la racine du projet :

    backend/.venv/Scripts/python.exe -m backend.docmodele "documents/Circulaire.docx"
"""
import copy
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

from .settings import settings

# Nom et emplacement par défaut du modèle. `modeles/` est ignoré par git : le
# logo appartient au lycée et n'a rien à faire ni dans le dépôt ni dans l'.exe.
SOUS_DOSSIER_MODELES = "modeles"
NOM_MODELE = "modele-etablissement.docx"

# Titre écrit dans les propriétés du fichier produit, à la place de l'auteur
# réel du document source (RGPD : le nom de session de l'agent qui a rédigé la
# circulaire ne doit pas se propager dans tous les courriers générés).
TITRE_MODELE = "Modèle de l'établissement"

# Repli quand AUCUN en-tête de page ne porte d'image : on conserve alors les
# premiers paragraphes du corps jusqu'à celui qui contient les images (inclus),
# vidés de leur texte. C'est le cas « ours dans le corps ».
_ENTETE_DANS_LE_CORPS = 6          # nb max de paragraphes de tête inspectés


def _dossier_application() -> Path:
    """Dossier de l'application — celui de l'exécutable en mode gelé.

    Même calcul que dans `ocr.py` : `sys._MEIPASS` ne conviendrait pas, c'est un
    dossier temporaire recréé à chaque lancement, alors que le modèle doit vivre
    à côté de l'.exe (clé USB) et survivre au redémarrage.
    """
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent.parent


def chemin_modele() -> Path:
    """Chemin du modèle : réglage `docgen_template_path`, sinon emplacement par défaut."""
    try:
        regle = str(settings.get().get("docgen_template_path") or "").strip()
    except Exception:
        regle = ""
    if regle:
        candidat = Path(regle)
        if candidat.is_dir():
            candidat = candidat / NOM_MODELE
        return candidat
    return _dossier_application() / SOUS_DOSSIER_MODELES / NOM_MODELE


# ---------- Inspection ----------
def _paragraphe_a_une_image(p) -> bool:
    return len(p.findall(".//" + qn("w:drawing"))) > 0


def _texte_de(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def _partie_referencee(doc, sectPr, tag: str):
    """Partie en-tête/pied explicitement référencée par un sectPr, sinon None.

    On lit la référence dans le XML plutôt que d'utiliser `section.header` :
    l'API de python-docx CRÉE un en-tête vide quand la section n'en définit pas,
    ce qui ajouterait au modèle une partie inutile — et, pour une section liée à
    la précédente, ferait perdre le lien.
    """
    for ref in sectPr.findall(qn(tag)):
        rid = ref.get(qn("r:id"))
        if not rid:
            continue
        try:
            return doc.part.related_parts[rid]
        except KeyError:
            continue
    return None


def _section_porteuse(doc):
    """Première section dont l'en-tête de page, explicitement défini, porte une image."""
    for section in doc.sections:
        partie = _partie_referencee(doc, section._sectPr, "w:headerReference")
        if partie is None:
            continue
        if len(partie.element.findall(".//" + qn("w:drawing"))) > 0:
            return section
    return None


def inspecter(chemin: Path) -> dict:
    """Décrit un .docx : images d'en-tête/pied, styles, marges, texte résiduel.

    Sert à la fois au contrôle après fabrication et à l'affichage de l'état dans
    les Paramètres. Ne lève jamais : renvoie `{"lisible": False}` en cas d'échec.
    """
    try:
        doc = Document(str(chemin))
    except Exception as e:
        return {"lisible": False, "erreur": str(e)}
    section = doc.sections[0]
    entete = _partie_referencee(doc, section._sectPr, "w:headerReference")
    pied = _partie_referencee(doc, section._sectPr, "w:footerReference")
    entete = entete.element if entete is not None else None
    pied = pied.element if pied is not None else None
    styles = {s.name for s in doc.styles if s.name}
    return {
        "lisible": True,
        "sections": len(doc.sections),
        "images_entete": 0 if entete is None else len(entete.findall(".//" + qn("w:drawing"))),
        "images_pied": 0 if pied is None else len(pied.findall(".//" + qn("w:drawing"))),
        "images_corps": len(doc.element.body.findall(".//" + qn("w:drawing"))),
        "paragraphes": len(doc.paragraphs),
        "texte_corps": len(_texte_de(doc.element.body).strip()),
        "texte_entete": 0 if entete is None else len(_texte_de(entete).strip()),
        "texte_pied": 0 if pied is None else len(_texte_de(pied).strip()),
        "police": doc.styles["Normal"].font.name,
        "styles": sorted(styles),
        "marges_cm": {
            "haut": round(section.top_margin.cm, 2),
            "bas": round(section.bottom_margin.cm, 2),
            "gauche": round(section.left_margin.cm, 2),
            "droite": round(section.right_margin.cm, 2),
        },
        "auteur": (doc.core_properties.author or ""),
    }


def etat() -> dict:
    """État du modèle, affiché dans Paramètres → Documents (motif repris d'OCR)."""
    chemin = chemin_modele()
    if not chemin.is_file():
        return {
            "disponible": False,
            "chemin": str(chemin),
            "message": "Aucun modèle de l'établissement : les documents seront créés "
                       "sans le logo ni l'en-tête du lycée.",
        }
    infos = inspecter(chemin)
    if not infos.get("lisible"):
        return {
            "disponible": False,
            "chemin": str(chemin),
            "message": f"Modèle illisible ({infos.get('erreur', 'erreur inconnue')}).",
        }
    images = infos["images_entete"] + infos["images_corps"]
    return {
        "disponible": True,
        "chemin": str(chemin),
        "message": f"Modèle de l'établissement prêt ({images} image(s) d'en-tête, "
                   f"police {infos.get('police') or 'par défaut'}).",
        "details": infos,
    }


# ---------- Fabrication ----------
def _vider_runs_de_texte(paragraphe) -> None:
    """Retire les runs qui portent du texte, garde ceux qui portent une image."""
    for run in list(paragraphe.findall(qn("w:r"))):
        if run.find(qn("w:drawing")) is not None:
            for t in list(run.findall(qn("w:t"))):
                run.remove(t)
            continue
        paragraphe.remove(run)


def _paragraphes_de_tete(doc) -> list:
    """Repli « ours dans le corps » : paragraphes de tête jusqu'aux images incluses."""
    gardes = []
    for p in doc.element.body.findall(qn("w:p"))[:_ENTETE_DANS_LE_CORPS]:
        copie = copy.deepcopy(p)
        _vider_runs_de_texte(copie)
        # Un sectPr accroché à ce paragraphe appartient à la section précédente :
        # il n'a rien à faire dans un modèle à section unique.
        pPr = copie.find(qn("w:pPr"))
        if pPr is not None:
            for sect in list(pPr.findall(qn("w:sectPr"))):
                pPr.remove(sect)
        gardes.append(copie)
        if _paragraphe_a_une_image(p):
            return gardes
    return []


def _purger_metadonnees(doc) -> None:
    """RGPD : efface l'auteur et les mentions du document source."""
    props = doc.core_properties
    props.author = ""
    props.last_modified_by = ""
    props.title = TITRE_MODELE
    props.subject = ""
    props.keywords = ""
    props.comments = ""
    props.category = ""


def construire_modele(source: Path, destination: Optional[Path] = None) -> dict:
    """Fabrique le modèle depuis `source` et l'écrit dans `destination`.

    Renvoie le compte rendu de l'inspection du fichier produit.
    Lève `ValueError` si la source est absente ou illisible.
    """
    source = Path(source)
    if not source.is_file():
        raise ValueError(f"Document source introuvable : {source}")
    if source.suffix.lower() != ".docx":
        raise ValueError("Le document source doit être un fichier Word (.docx).")
    destination = Path(destination) if destination else chemin_modele()

    try:
        doc = Document(str(source))
    except Exception as e:
        raise ValueError(f"Document source illisible : {e}")

    porteuse = _section_porteuse(doc)
    de_tete = [] if porteuse is not None else _paragraphes_de_tete(doc)
    section = porteuse if porteuse is not None else doc.sections[0]
    sectPr = copy.deepcopy(section._sectPr)
    # Section unique : le modèle ne doit pas hériter d'un saut de page.
    for t in list(sectPr.findall(qn("w:type"))):
        sectPr.remove(t)

    body = doc.element.body
    body.clear_content()
    ancien = body.find(qn("w:sectPr"))
    if ancien is not None:
        body.remove(ancien)
    for p in de_tete:
        body.append(p)
    body.append(sectPr)

    _purger_metadonnees(doc)
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(destination))

    infos = inspecter(destination)
    infos["source"] = source.name
    infos["chemin"] = str(destination)
    return infos


def _principal(argv: list) -> int:
    if len(argv) < 2:
        print("Usage : python -m backend.docmodele <document-source.docx> [destination.docx]")
        return 2
    dest = Path(argv[2]) if len(argv) > 2 else None
    try:
        infos = construire_modele(Path(argv[1]), dest)
    except ValueError as e:
        print(f"Echec : {e}")
        return 1
    print(f"Modele ecrit : {infos['chemin']}")
    print(f"  images en-tete : {infos['images_entete']} / pied : {infos['images_pied']}")
    print(f"  texte residuel : {infos['texte_corps']} caractere(s) dans le corps")
    print(f"  police Normal  : {infos['police']}")
    print(f"  marges (cm)    : {infos['marges_cm']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(_principal(sys.argv))
